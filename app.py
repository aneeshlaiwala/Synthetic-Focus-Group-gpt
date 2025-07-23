import streamlit as st
import base64
from io import BytesIO
from docx import Document
from docx.shared import Pt
from PyPDF2 import PdfReader
import docx2txt

# Page setup
st.set_page_config(page_title="FGD Generator", layout="wide")
st.markdown("<h1 style='text-align: center; color: #6A1B9A;'>🎤 Synthetic Focus Group Discussion Generator</h1>", unsafe_allow_html=True)
st.markdown("---")

# Sidebar - AI Engine Settings
st.sidebar.header("🧠 AI Engine Settings")
ai_tool = st.sidebar.selectbox("Select AI Tool", ["OpenAI", "Anthropic", "Google Gemini", "Mistral", "HuggingFace"])
model_options = {
    "OpenAI": ["gpt-4", "gpt-4o"],
    "Anthropic": ["claude-3-opus"],
    "Google Gemini": ["gemini-1.5-pro"],
    "Mistral": ["mixtral-8x7b"],
    "HuggingFace": ["llama3-8b"]
}
model = st.sidebar.selectbox("Select Model", model_options[ai_tool])
api_key = st.sidebar.text_input("Enter API Key", type="password")

# Main Section - Discussion Setup
st.subheader("📋 Discussion Setup")

# Participant Composition Inputs
st.markdown("### 🗂️ Participant Composition")
col1, col2, col3, col4 = st.columns([1, 1, 1, 1])
with col1:
    # 1. Total Participants input
    num_participants = st.number_input("👥 Total Participants", min_value=2, max_value=20, value=6, key="num_participants")
with col2:
    # 2. Male Participants input (max constrained by total)
    male_count = st.number_input("🧔‍♂️ Male Participants", min_value=0, max_value=st.session_state.num_participants, value=2, key="male_count")
with col3:
    # 2. Female Participants input (max constrained by remaining slots)
    female_max = max(0, st.session_state.num_participants - st.session_state.male_count)  # remaining after male
    female_count = st.number_input("👩 Female Participants", min_value=0, max_value=female_max, value=2, key="female_count")
with col4:
    # 3. Non-Binary Participants auto-calculated display
    nb_count = st.session_state.num_participants - st.session_state.male_count - st.session_state.female_count
    st.markdown(
        f"<div style='margin-top: 5px; font-size: 16px;'>🧑‍🎤 <span style='color: #9C27B0; font-weight: bold;'>Non-Binary Participants: {nb_count}</span> <em style='color: grey;'>(auto-calculated)</em></div>",
        unsafe_allow_html=True
    )

# 4. Overshoot check: If Male + Female > Total, show warning and auto-limit
if st.session_state.male_count + st.session_state.female_count > st.session_state.num_participants:
    # Auto-limit values to maintain the total
    st.session_state.male_count = min(st.session_state.male_count, st.session_state.num_participants)
    st.session_state.female_count = max(0, st.session_state.num_participants - st.session_state.male_count)
    st.warning("⚠️ **Male and Female participants exceed Total!** Values have been adjusted to fit the total.")

# Group Profile & Session Settings
st.markdown("### 🌍 Group Profile & Session Settings")
col_gp1, col_gp2, col_gp3 = st.columns([1, 1, 1])
with col_gp1:
    age_range = st.text_input("📅 Age Range", "25–35", key="age_range")
with col_gp2:
    location = st.text_input("📍 Location (City, Country)", "Mumbai, India", key="location")
with col_gp3:
    language = st.multiselect("🗣️ Language(s)", ["English", "Hindi", "Hinglish", "French"], default=["Hinglish"], key="language")

mode = st.selectbox("🧭 Mode of Discussion", ["Online", "Offline"], key="mode")
duration_min = st.slider("⏱️ Duration (minutes)", 10, 120, 60, key="duration_min")

st.markdown("### 🧩 Topic & Demographics")
topic = st.text_input("📌 Discussion Topic", "Electric Vehicles in Tier 2 Indian Cities", key="topic")
demo_profile = st.text_area("👤 Demographic Profile", "Mid-income professionals, mix of adopters and skeptics", key="demo_profile")

st.markdown("### 📄 Study Objective")
study_objective = st.text_area("✍️ Write Study Objective", "", key="study_objective")
uploaded_file = st.file_uploader("📎 Or Upload a File (.pdf, .docx, .txt)", type=["pdf", "docx", "txt"])

# Generate Prompt Button
if st.button("Generate Prompt"):
    # Combine uploaded file content with study_objective, if any
    full_objective = study_objective
    if uploaded_file:
        file_text = ""
        try:
            if uploaded_file.type == "application/pdf":
                reader = PdfReader(uploaded_file)
                file_text = "\n".join(page.extract_text() for page in reader.pages)
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                file_text = docx2txt.process(uploaded_file)
            else:
                file_text = uploaded_file.read().decode("utf-8")
        except Exception:
            file_text = ""
        if file_text:
            # Append file content to the written objective with spacing
            full_objective = (study_objective + "\n\n" + file_text).strip()
    # Word count estimate based on duration
    word_count = int(st.session_state.duration_min * 140)
    # Construct the prompt string with all inputs
    prompt = f"""You are an expert qualitative researcher and focus group moderator.

Simulate a highly realistic and detailed synthetic focus group discussion (FGD) transcript.

--- PARTICIPANT PROFILE ---
- Total Participants: {st.session_state.num_participants} ({st.session_state.male_count} male, {st.session_state.female_count} female, {st.session_state.num_participants - st.session_state.male_count - st.session_state.female_count} non-binary)
- Age Range: {age_range}
- Location: {location}
- Demographic: {demo_profile}
- Languages spoken: {', '.join(language)}
- Mode: {mode}
- Duration: {duration_min} minutes (~{word_count} words)
- Discussion Topic: {topic}
- Study Objective: {full_objective}

--- STRUCTURE TO FOLLOW ---
1. Opening (welcome, intros, ground rules, consent)
2. Warm-up (icebreaker, general thoughts)
3. Core Discussion (key Qs, probing, opposing views)
4. Closing (summary, final comments, thank you)

--- MODERATOR BEHAVIOR ---
- Remain neutral and professional
- Use natural transitions
- Encourage quieter participants
- Handle dominant voices politely
- Use real participant names typical of {location}
- Allow grammatical imperfections and local dialects

--- RESEARCH REQUIREMENT ---
Thoroughly research the discussion topic using multiple reputable sources specific to the location. Reflect realistic, localized views based on what people are actually discussing or facing in this region. Do not rely on generic or single-source assumptions.

Format the output as a readable transcript with MODERATOR and PARTICIPANT names."""
    # Show the generated prompt in a text area for review/editing
    st.markdown("### ✏️ Editable Prompt")
    final_prompt = st.text_area("You may revise the prompt before submission", prompt, height=400)
    st.session_state["final_prompt"] = final_prompt  # store prompt in session state

    # Transcript Generation (Simulated Demo)
    if st.button("🚀 Generate Transcript (Simulated Demo)"):
        demo_transcript = f"""MODERATOR: Welcome everyone. Let's introduce ourselves briefly.

RAHUL (Male, 28): I'm Rahul from Dadar. I work as a software engineer. Curious about EVs.

SNEHA (Female, 30): Hi! I'm Sneha. I’ve started thinking about switching to an EV recently.

AADI (Non-Binary, 26): Hello! I live in Thane and ride a scooter. Wondering if an EV makes sense...

...

MODERATOR: Thank you all for your views today. Your inputs on "{topic}" were incredibly helpful.""".strip()
        st.markdown("### 📄 Synthetic Transcript")
        st.text_area("Transcript", demo_transcript, height=300)
        st.download_button("⬇️ Download TXT", data=demo_transcript, file_name="transcript.txt", mime="text/plain")

        # Prepare a DOCX file for download
        doc = Document()
        doc.add_heading("Focus Group Discussion Transcript", 0)
        for line in demo_transcript.split("\n"):
            para = doc.add_paragraph(line.strip())
            if "MODERATOR:" in line:
                para.runs[0].bold = True
            para.style.font.size = Pt(11)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        st.download_button("⬇️ Download DOCX", data=buffer, file_name="transcript.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
